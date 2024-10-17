import streamlit as st
import pandas as pd
import requests
from io import BytesIO
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from docx import Document
from docx.shared import Pt
import pdfplumber

def extract_text_from_pdf(pdf_file):
    text = ""
    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            text += page.extract_text()
    return text

# Function to extract text from a CSV file
def extract_text_from_csv(csv_file):
    df = pd.read_csv(csv_file)
    return df.to_string()

# Function to extract text from an Excel file
def extract_text_from_excel(excel_file):
    df = pd.read_excel(excel_file)
    return df.to_string()

# Function to preprocess text
def preprocess_text(text):
    return text.lower()

# Function to query the Gemini API
def query_gemini_api(user_query, context_data, api_key):
    url = 'https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash-latest:generateContent'
    headers = {'Content-Type': 'application/json'}
    payload = {
        "contents": [
            {"parts": [{"text": f"Use only the following context: {context_data} \n\n{user_query}"}]}
        ]
    }
    
    response = requests.post(f"{url}?key={api_key}", headers=headers, json=payload)
    
    if response.status_code == 429:
        return {"error": "API quota exceeded. Please try again later."}
    
    if response.status_code != 200:
        return {"error": response.text}
    
    return response.json()

# Function to create a PDF for chat history using reportlab with A4 size
def create_chat_history_pdf(chat_history):
    pdf_output = BytesIO()
    c = canvas.Canvas(pdf_output, pagesize=A4)
    width, height = A4

    y_position = height - 50  # Start position at the top of the page

    def wrap_text(text, max_words_per_line=12):
        words = text.split()
        lines = []
        for i in range(0, len(words), max_words_per_line):
            line = " ".join(words[i:i+max_words_per_line])
            lines.append(line)
        return lines

    for query, response in chat_history:
        # Add user query
        c.setFont("Helvetica-Bold", 12)
        c.setFillColor(colors.red)
        c.drawString(50, y_position, f"You: {query}")
        y_position -= 20

        # Add chatbot response
        c.setFont("Helvetica", 12)
        c.setFillColor(colors.black)

        # Split response into lines with max 12 words per line
        wrapped_response_lines = wrap_text(response, max_words_per_line=10)
        for line in wrapped_response_lines:
            if y_position < 40:  # If we are near the bottom of the page, create a new page
                c.showPage()
                y_position = height - 50
            c.drawString(50, y_position, f"Chatbot: {line}")
            y_position -= 20

        # Add extra space between conversations
        y_position -= 10

        if y_position < 50:
            c.showPage()
            y_position = height - 50

    c.save()
    pdf_output.seek(0)
    return pdf_output

# Function to create Excel file for chat history
def create_chat_history_excel(chat_history):
    output = BytesIO()
    writer = pd.ExcelWriter(output, engine='xlsxwriter')
    
    data = {'Query': [item[0] for item in chat_history],
            'Response': [item[1] for item in chat_history]}
    
    df = pd.DataFrame(data)
    df.to_excel(writer, index=False, sheet_name='Chat History')

    writer.save()
    output.seek(0)
    return output

# Function to create DOCX file for chat history
def create_chat_history_docx(chat_history):
    doc = Document()
    for query, response in chat_history:
        doc.add_heading('You:', level=2)
        doc.add_paragraph(query)
        doc.add_heading('Chatbot:', level=2)
        doc.add_paragraph(response)
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# Function to create CSV file for chat history
def create_chat_history_csv(chat_history):
    output = BytesIO()
    data = {'Query': [item[0] for item in chat_history],
            'Response': [item[1] for item in chat_history]}
    
    df = pd.DataFrame(data)
    df.to_csv(output, index=False)
    output.seek(0)
    return output

# Streamlit app
def main():
    st.title("Chat with your PDF, CSV, DOCX, and Excel Files \n by: Mohammad R. Abu Ayyash")

    api_key = st.text_input("Enter Gemini API Key(model: gemini 1.5 flash):", type="password")

    uploaded_file = st.file_uploader("Upload a PDF, CSV, DOCX, or Excel file", type=["pdf", "csv", "docx", "xlsx"])

    if 'extracted_text' not in st.session_state:
        st.session_state.extracted_text = ""
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []
    if 'user_query' not in st.session_state:
        st.session_state.user_query = ""

    if uploaded_file is not None:
        file_type = uploaded_file.type

        if file_type == "application/pdf":
            pdf_text = extract_text_from_pdf(uploaded_file)
            st.write("PDF text extracted successfully!")
            st.session_state.extracted_text = preprocess_text(pdf_text)

        elif file_type == "text/csv":
            csv_text = extract_text_from_csv(uploaded_file)
            st.write("CSV text extracted successfully!")
            st.session_state.extracted_text = preprocess_text(csv_text)
        
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(uploaded_file)
            doc_text = "\n".join([para.text for para in doc.paragraphs])
            st.write("DOCX text extracted successfully!")
            st.session_state.extracted_text = preprocess_text(doc_text)

        elif file_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
            excel_text = extract_text_from_excel(uploaded_file)
            st.write("Excel text extracted successfully!")
            st.session_state.extracted_text = preprocess_text(excel_text)

    user_query = st.text_input("Ask a question based on the content:", value=st.session_state.user_query)

    if user_query and api_key:
        if st.session_state.extracted_text:
            response = query_gemini_api(user_query, st.session_state.extracted_text, api_key)

            if "error" in response:
                st.write("Error:", response["error"])
            else:
                chatbot_response = response.get("candidates", [{}])[0].get("content", {}).get("parts", [{}])[0].get("text", "No response")
                st.write("Chatbot Response:")
                st.write(chatbot_response)

                st.session_state.chat_history.append((user_query, chatbot_response))
                st.session_state.user_query = ""  # Clear the input after submitting
        else:
            st.write("Please upload a file to extract text before asking questions.")

    if st.session_state.chat_history:
        st.write("### Chat History:")
        chat_display = ""
        for query, response in st.session_state.chat_history:
            chat_display += f"**You:** {query}\n**Chatbot:** {response}\n\n"
        
        # Display the chat history in a markdown block for easy copy-pasting
        st.markdown(f"```markdown\n{chat_display}\n```")

        # Add a selection box to choose the file format
        file_format = st.selectbox(
            "Select the format to download the chat history:",
            ("PDF", "DOCX", "CSV")
        )

        if file_format == "PDF":
            pdf_file = create_chat_history_pdf(st.session_state.chat_history)
            st.download_button(
                label="Download Chat History as PDF",
                data=pdf_file,
                file_name="chat_history.pdf",
                mime="application/pdf"
            )
        elif file_format == "DOCX":
            docx_file = create_chat_history_docx(st.session_state.chat_history)
            st.download_button(
                label="Download Chat History as DOCX",
                data=docx_file,
                file_name="chat_history.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        elif file_format == "CSV":
            csv_file = create_chat_history_csv(st.session_state.chat_history)
            st.download_button(
                label="Download Chat History as CSV",
                data=csv_file,
                file_name="chat_history.csv",
                mime="text/csv"
            )

if __name__ == "__main__":
    main()
